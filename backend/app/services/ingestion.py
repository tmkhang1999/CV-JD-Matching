# app/services/ingestion.py
from pathlib import Path
from fastapi import UploadFile
import uuid
import fitz  # PyMuPDF
import docx  # python-docx
import re

UPLOAD_DIR = Path("data/uploads")


def save_upload_file(file: UploadFile) -> Path:
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    ext = Path(file.filename).suffix
    filepath = UPLOAD_DIR / f"{uuid.uuid4()}{ext}"
    with filepath.open("wb") as f:
        f.write(file.file.read())
    return filepath


def clean_text(text: str) -> str:
    """Clean text by removing NUL characters and other problematic characters."""
    if not text:
        return ""

    # Remove NUL characters that cause PostgreSQL errors
    text = text.replace("\x00", "")

    # Remove other control characters except newlines, tabs, and carriage returns
    text = re.sub(r"[\x01-\x08\x0B\x0C\x0E-\x1F\x7F]", "", text)

    # Normalize whitespace
    text = re.sub(r"\s+", " ", text)
    text = text.strip()

    return text


def extract_text_from_pdf(path: Path) -> str:
    doc = fitz.open(path)
    texts = []
    for page in doc:
        page_text = page.get_text()
        texts.append(clean_text(page_text))
    return "\n".join(texts)


def extract_text_from_docx(path: Path) -> str:
    """Extract text from DOCX files including paragraphs, tables, headers, and footers."""
    doc = docx.Document(str(path))
    texts = []

    # Extract from paragraphs
    for para in doc.paragraphs:
        text = clean_text(para.text)
        if text:
            texts.append(text)

    # Extract from tables (many CVs use table layouts)
    for table in doc.tables:
        for row in table.rows:
            row_texts = []
            for cell in row.cells:
                cell_text = clean_text(cell.text)
                if cell_text and cell_text not in row_texts:  # Avoid duplicates from merged cells
                    row_texts.append(cell_text)
            if row_texts:
                texts.append(" | ".join(row_texts))

    # Extract from headers
    for section in doc.sections:
        header = section.header
        if header:
            for para in header.paragraphs:
                text = clean_text(para.text)
                if text:
                    texts.append(text)

        # Extract from footers
        footer = section.footer
        if footer:
            for para in footer.paragraphs:
                text = clean_text(para.text)
                if text:
                    texts.append(text)

    return "\n".join(texts)


def extract_raw_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(path)
    else:
        # Fallback: try to read as plain text
        text = path.read_text(errors="ignore")
        return clean_text(text)
